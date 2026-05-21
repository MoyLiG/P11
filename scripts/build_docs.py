"""Generation des livrables documentaires .docx et .pptx.

Convertit :
- docs/rapport_technique.md  ->  docs/rapport_technique.docx
- docs/deroule_projet.md     ->  docs/deroule_projet.docx
Puis genere docs/presentation.pptx (12 slides).

Dependances : python-docx, python-pptx (a installer separement, hors
requirements.txt principal puisque uniquement utiles pour generer la
documentation).

Usage :
    pip install python-docx==1.1.2 python-pptx==1.0.2
    python scripts/build_docs.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

# Le builder .pptx est dans son propre module (richesse de schemas).
sys_path_marker = None
try:
    from build_pptx import build as build_pptx_rich
except ImportError:
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from build_pptx import build as build_pptx_rich

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"


# =====================================================================
# Conversion Markdown -> .docx (limite a notre dialecte : titres, paragraphes,
# listes, tableaux pipe, blocs de code triples backticks)
# =====================================================================

def md_to_docx(md_path: Path, out_path: Path, title: str) -> None:
    """Convertit un Markdown en .docx avec une mise en forme propre."""
    doc = Document()

    # Marges + style de base
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titre du document (avant le contenu)
    title_p = doc.add_heading(title, level=0)
    title_p.alignment = 1  # CENTER

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        # On filtre les lignes de separation (---)
        rows = [
            r for r in table_buffer
            if not all(re.fullmatch(r"\s*:?-+:?\s*", c) for c in r)
        ]
        if not rows:
            table_buffer = []
            return
        ncols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncols)
        table.style = "Light Grid Accent 1"
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell_text = row[ci] if ci < len(row) else ""
                cell = table.cell(ri, ci)
                cell.text = cell_text
                if ri == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
        doc.add_paragraph("")
        table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Bloc de code
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Detection de tableau (ligne pipe)
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_buffer.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # Titres
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = min(len(m.group(1)), 5)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Separateur horizontal
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # Listes a puces
        m_bul = re.match(r"^\s*[-*]\s+(.+)$", line)
        if m_bul:
            doc.add_paragraph(_strip_md_inline(m_bul.group(1)), style="List Bullet")
            i += 1
            continue

        # Listes numerotees
        m_num = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if m_num:
            doc.add_paragraph(_strip_md_inline(m_num.group(1)), style="List Number")
            i += 1
            continue

        # Citation
        m_q = re.match(r"^\s*>\s*(.*)$", line)
        if m_q:
            p = doc.add_paragraph(_strip_md_inline(m_q.group(1)))
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Ligne vide
        if not line.strip():
            i += 1
            continue

        # Paragraphe normal
        _add_paragraph_with_inline(doc, line)
        i += 1

    flush_table()
    doc.save(str(out_path))
    print(f"OK: {out_path}")


_STAR_PATTERN = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`")


def _strip_md_inline(text: str) -> str:
    """Enleve simplement les marqueurs de mise en forme inline."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _add_paragraph_with_inline(doc: Document, text: str) -> None:
    """Ajoute un paragraphe en preservant gras / italique / inline-code."""
    p = doc.add_paragraph()
    pos = 0
    for match in _STAR_PATTERN.finditer(text):
        if match.start() > pos:
            p.add_run(text[pos:match.start()])
        bold, italic, code = match.group(1), match.group(2), match.group(3)
        if bold:
            run = p.add_run(bold)
            run.bold = True
        elif italic:
            run = p.add_run(italic)
            run.italic = True
        elif code:
            run = p.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        p.add_run(text[pos:])


# =====================================================================
# Generation .pptx : delegue au module build_pptx.py
# (presentation riche en schemas vectoriels - dark theme)
# Le contenu detaille des slides est dans scripts/build_pptx.py.
# =====================================================================

# Conservation desactivee : ancienne version texte-only (gardee pour reference).
_LEGACY_SLIDES_DISABLED = [
    {
        "title": "POC RAG Puls-Events",
        "subtitle": "Chatbot d'evenements culturels - Pays de la Loire\n"
                    "OpenClassrooms - Data Engineer - Projet 11\n"
                    "Morgan Moy - mai 2026",
    },
    {
        "title": "Contexte & objectif",
        "bullets": [
            "Puls-Events : plateforme de decouverte d'evenements culturels.",
            "Besoin metier : chatbot capable de recommander en langage naturel.",
            "Objectif POC : valider la faisabilite technique avant deploiement.",
            "Stack imposee : LangChain + Mistral + FAISS.",
            "Source : Open Agenda (dataset Opendatasoft, 1,1 M evenements au catalogue).",
        ],
    },
    {
        "title": "Perimetre du POC",
        "bullets": [
            "Region : Pays de la Loire (configurable via config.yaml).",
            "Fenetre temporelle : evenements de moins d'un an (firstdate_begin >= today - 365j).",
            "Volume cible : ~3 000 a 6 000 evenements actifs.",
            "Pourquoi Pays de la Loire : volume suffisant + diversite (Nantes, Hellfest, Folle Journee, Voyage a Nantes).",
        ],
    },
    {
        "title": "Architecture",
        "bullets": [
            "1. Ingestion : Opendatasoft v2.1 + filtres ODSQL + retry tenacity.",
            "2. Preprocessing : clean HTML, dedupe (uid), filtre region/date.",
            "3. Indexation : chunking 1200/80 + mistral-embed + FAISS IndexFlatL2 + CacheBackedEmbeddings.",
            "4. RAG : retriever MMR k=4 (fetch_k=12) + prompt FR balisé + mistral-small-latest.",
            "Demo : CLI Python + app Streamlit, conteneurisée via Docker.",
        ],
    },
    {
        "title": "Reproductibilite : Docker",
        "bullets": [
            "Stack Python data = Linux-first (LangChain, Mistral SDK, FAISS, Streamlit).",
            "Dockerfile : python:3.12-slim, user non-root, healthcheck, ~600 Mo.",
            "docker-compose.yml : volumes persistants data/, limites 2 GB / 2 CPU.",
            "Demo evaluateur : 'docker compose run --rm rag python scripts/pipeline.py' puis 'docker compose up rag'.",
            "Trois paths documentes : WSL2/Linux (recommande), Docker (reproductible), PowerShell (alternatif).",
        ],
    },
    {
        "title": "Choix techniques cles",
        "bullets": [
            "FAISS IndexFlatL2 : exact, suffisant pour < 100k chunks.",
            "Chunking RecursiveCharacterTextSplitter : preserve la structure FR.",
            "MMR (Maximal Marginal Relevance) : evite la redondance dans le top-k.",
            "Prompt systeme : citation obligatoire + refus explicite hors-scope.",
            "mistral-small-latest : T=0.2 - rapport qualite/cout pour POC.",
        ],
    },
    {
        "title": "Pipeline de donnees",
        "bullets": [
            "scripts/01_fetch_data.py - telecharge -> data/raw/events.json.",
            "scripts/02_build_index.py - preprocess + vectorise -> data/vectorstore/.",
            "Pipeline complet via scripts/pipeline.py (idempotent, reproductible).",
            "Index reconstructible a la demande, conformement aux consignes.",
        ],
    },
    {
        "title": "Qualite des donnees - tests pytest",
        "bullets": [
            "test_data_freshness.py : 100% des events < 1 an.",
            "test_data_geography.py : 100% des events dans la region cible.",
            "test_preprocessing.py : clean_html, dedupe, build_document_text.",
            "test_vectorstore.py : split en chunks, validation cle API.",
            "Double check : filtre cote API + re-filtre Python verifie en test.",
        ],
    },
    {
        "title": "Demo live",
        "bullets": [
            "App Streamlit : chat input, sources cliquables vers Open Agenda.",
            "CLI Python : pour tests rapides et debugging.",
            "Sidebar : config courante (region, modele, top-k).",
            "Exemples a tester en live :",
            "   - 'Quels concerts a Nantes ?'",
            "   - 'Festival classique en Pays de la Loire ?'",
            "   - 'Que faire ce week-end a Saint-Nazaire ?'",
        ],
    },
    {
        "title": "Evaluation",
        "bullets": [
            "Jeu Q/R annote : 20 paires (5 categories + 3 hors-scope).",
            "Metriques : hit_rate@k, cosine similarity, LLM-as-judge (0-5).",
            "Cibles : hit_rate >= 70%, cosine >= 0.75, juge >= 3.5/5.",
            "Sortie : data/eval/results.csv + resume console.",
        ],
    },
    {
        "title": "Resultats POC",
        "bullets": [
            "Volume : ~5 000 chunks indexes apres preprocessing.",
            "Latence : ~1,5 s par requete (retriever 50 ms + LLM 1,4 s).",
            "Empreinte disque : ~25 Mo pour l'index FAISS.",
            "Cout : ~0,30 EUR par construction d'index (mistral-embed).",
            "Qualite : reponses pertinentes + refus correct hors-scope.",
        ],
    },
    {
        "title": "Limites du POC",
        "bullets": [
            "Filtre temporel sur firstdate_begin uniquement (variante possible : lastdate_end).",
            "IndexFlatL2 ne scale pas au-dela de 100k chunks.",
            "Pas d'historique conversationnel (consigne POC).",
            "Pas de re-ranking ni hybrid search (BM25).",
            "Pas de monitoring de drift.",
        ],
    },
    {
        "title": "Recommandations v1",
        "bullets": [
            "Index hybride : BM25 + dense (Reciprocal Rank Fusion).",
            "Re-ranker : bge-reranker-v2-m3 entre retriever et LLM.",
            "Modele plus large (mistral-large) pour la generation finale.",
            "Indexation incrementale (champ updatedat).",
            "Deploiement : image Docker -> ECS / Cloud Run / Kubernetes.",
            "Observabilite : LangSmith / Phoenix + RAGAS en CI nocturne.",
            "KPI produit : CTR sources, taux de refus, satisfaction utilisateur.",
        ],
    },
    {
        "title": "Conclusion",
        "bullets": [
            "Le POC valide la stack LangChain + Mistral + FAISS.",
            "Le pipeline est entierement reproductible.",
            "La qualite est suffisante pour passer en v1.",
            "Les couts operationnels sont maitrisables.",
            "Roadmap v1 : 6-12 semaines pour passer du POC a la production.",
            "Merci - questions ?",
        ],
    },
]


def build_pptx(out_path: Path) -> None:
    """Genere la presentation PowerPoint via le builder schema-first."""
    build_pptx_rich(out_path)


# =====================================================================
# Entree
# =====================================================================

def main() -> int:
    md_to_docx(
        DOCS / "rapport_technique.md",
        DOCS / "rapport_technique.docx",
        title="Rapport technique - POC RAG Puls-Events",
    )
    md_to_docx(
        DOCS / "deroule_projet.md",
        DOCS / "deroule_projet.docx",
        title="Deroule du projet - POC RAG Puls-Events",
    )
    build_pptx(DOCS / "presentation.pptx")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
